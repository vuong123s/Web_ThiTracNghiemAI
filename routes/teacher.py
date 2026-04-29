import io
import json
import logging
import os
import random
import re
import time
from datetime import datetime
from functools import wraps
from urllib import error as urllib_error
from urllib import parse as urllib_parse
from urllib import request as urllib_request
from urllib.parse import quote

from flask import (
    Blueprint,
    current_app,
    flash,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    url_for,
)
from flask_login import current_user, login_required
from openpyxl import Workbook, load_workbook
from werkzeug.utils import secure_filename

from forms import (
    AIGenerateQuestionsForm,
    ClassroomForm,
    DocumentUploadForm,
    EssayGradingForm,
    ImportBankForm,
    QuestionBankForm,
    QuizBuilderForm,
)
from models import (
    Answer,
    BankOption,
    ClassDocument,
    Classroom,
    Option,
    Question,
    QuestionBank,
    Quiz,
    Submission,
    UserClass,
    db,
)

logger = logging.getLogger(__name__)
teacher_bp = Blueprint("teacher", __name__)


def teacher_required(f):
    @wraps(f)
    @login_required
    def decorated_function(*args, **kwargs):
        if not current_user.is_teacher():
            flash("Chỉ giáo viên mới có quyền truy cập trang này.", "danger")
            return redirect(url_for("dashboard"))
        return f(*args, **kwargs)

    return decorated_function


def _parse_datetime(raw_value: str | None):
    if not raw_value:
        return None
    value = raw_value.strip()
    if not value:
        return None
    for fmt in ("%Y-%m-%d %H:%M", "%Y-%m-%dT%H:%M"):
        try:
            return datetime.strptime(value, fmt)
        except ValueError:
            continue
    return None


def _normalize_difficulty(value: str | None, default: str = "medium") -> str:
    if not value:
        return default
    normalized = value.strip().lower()
    if normalized in {"easy", "medium", "hard"}:
        return normalized
    return default


def _extract_correct_letters(raw_value: str | None) -> set[str]:
    if not raw_value:
        return set()
    letters = set()
    for chunk in raw_value.split(","):
        value = chunk.strip().upper()
        if value in {"A", "B", "C", "D"}:
            letters.add(value)
    return letters


def _get_ai_provider_status() -> dict:
    """
    Kiểm tra trạng thái của các AI providers.
    Trả về dict với status của OpenAI và Gemini.
    """
    openai_key = current_app.config.get("OPENAI_API_KEY")
    gemini_key = current_app.config.get("GEMINI_API_KEY")
    
    return {
        "openai": {
            "ready": bool(openai_key and openai_key.strip()),
            "status": "OpenAI đã sẵn sàng" if (openai_key and openai_key.strip()) else "Chưa cấu hình OPENAI_API_KEY trong .env",
        },
        "gemini": {
            "ready": bool(gemini_key and gemini_key.strip()),
            "status": "Gemini đã sẵn sàng" if (gemini_key and gemini_key.strip()) else "Chưa cấu hình GEMINI_API_KEY trong .env",
        },
    }


def _is_retryable_error(error_code: int, error_msg: str) -> bool:
    """
    Kiểm tra lỗi có thể retry không.
    - 429 (Rate limit): Có thể retry sau một thời gian
    - 500, 502, 503 (Server error): Có thể retry
    - 401, 404, 403: Không retry (lỗi cấu hình)
    - insufficient_quota: Không retry (cần nạp credit)
    """
    # Lỗi không thể retry
    if error_code in {400, 401, 403, 404}:
        return False
    
    # insufficient_quota - không retry
    if "insufficient_quota" in error_msg.lower() or "quota" in error_msg.lower():
        return False
    
    # Các lỗi có thể retry
    if error_code in {429, 500, 502, 503, 504}:
        return True
    
    # Lỗi mạng hoặc timeout
    if "timeout" in error_msg.lower() or "connection" in error_msg.lower():
        return True
    
    return False


def _call_ai_with_retry(
    provider: str,
    topic: str,
    question_count: int,
    difficulty: str,
    attempt: int = 1
) -> list[dict]:
    """
    Gọi AI provider với cơ chế retry và exponential backoff.
    Nếu provider fail và fallback enable, thử provider khác.
    
    Args:
        provider: "openai" hoặc "gemini"
        topic, question_count, difficulty: Tham số sinh câu hỏi
        attempt: Số lần thử hiện tại (dùng cho recursive retry)
    
    Returns:
        Danh sách câu hỏi từ AI
    
    Raises:
        RuntimeError: Nếu cả hai provider đều fail
    """
    max_retries = current_app.config.get("AI_MAX_RETRIES", 3)
    retry_delay = current_app.config.get("AI_RETRY_DELAY", 2)
    backoff_multiplier = current_app.config.get("AI_RETRY_BACKOFF", 2.0)
    enable_fallback = current_app.config.get("AI_ENABLE_FALLBACK", True)
    fallback_provider = current_app.config.get("AI_FALLBACK_PROVIDER", "gemini")
    
    try:
        if provider == "gemini":
            return _call_gemini_generate(topic, question_count, difficulty)
        else:
            return _call_openai_generate(topic, question_count, difficulty)
    
    except Exception as exc:
        error_msg = str(exc)
        error_code = None
        
        # Trích error code từ error message
        if "API lỗi (" in error_msg:
            match = re.search(r'API lỗi \((\d+)\)', error_msg)
            if match:
                error_code = int(match.group(1))
        
        logger.warning(f"Provider {provider} failed (attempt {attempt}/{max_retries}): {error_msg[:200]}")
        
        # Kiểm tra có thể retry không
        if attempt < max_retries and _is_retryable_error(error_code or 500, error_msg):
            # Exponential backoff
            delay = retry_delay * (backoff_multiplier ** (attempt - 1))
            logger.info(f"Retrying {provider} after {delay}s (attempt {attempt}/{max_retries})")
            time.sleep(delay)
            return _call_ai_with_retry(provider, topic, question_count, difficulty, attempt + 1)
        
        # Nếu không thể retry, thử fallback provider
        if enable_fallback and provider != fallback_provider:
            logger.info(f"Fallback từ {provider} sang {fallback_provider}")
            try:
                return _call_ai_with_retry(fallback_provider, topic, question_count, difficulty, attempt=1)
            except Exception as fallback_exc:
                # Cả hai provider đều fail
                combined_msg = f"{provider} failed: {error_msg[:150]}. Fallback ({fallback_provider}) failed: {str(fallback_exc)[:150]}"
                logger.error(f"Cả hai provider đều thất bại: {combined_msg}")
                raise RuntimeError(f"❌ Cả {provider} và {fallback_provider} đều không thể tạo câu hỏi. {combined_msg[:300]}")
        
        # Không thể retry hoặc fallback
        raise exc

def _sync_bank_options(bank_question: QuestionBank, form: QuestionBankForm) -> None:
    bank_question.options.delete()

    if bank_question.type == "essay":
        return

    option_payloads = [
        ("A", form.option_a.data),
        ("B", form.option_b.data),
        ("C", form.option_c.data),
        ("D", form.option_d.data),
    ]
    correct_letters = _extract_correct_letters(form.correct_answers.data)

    for idx, (letter, text) in enumerate(option_payloads, start=1):
        if not text or not text.strip():
            continue
        db.session.add(
            BankOption(
                question_bank_id=bank_question.id,
                text=text.strip(),
                is_correct=letter in correct_letters,
                order_index=idx,
            )
        )


def _import_from_excel(file_storage, teacher_id: int, default_category: str, default_difficulty: str):
    workbook = load_workbook(file_storage)
    sheet = workbook.active

    imported = 0
    errors = []

    for row_index, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=2):
        if not row or not row[0]:
            continue

        text = str(row[0]).strip()
        qtype = (str(row[1]).strip().lower() if len(row) > 1 and row[1] else "single")
        category = (str(row[2]).strip() if len(row) > 2 and row[2] else default_category) or None
        difficulty = _normalize_difficulty(
            str(row[3]).strip() if len(row) > 3 and row[3] else default_difficulty
        )
        tags = str(row[4]).strip() if len(row) > 4 and row[4] else None
        option_a = str(row[5]).strip() if len(row) > 5 and row[5] else ""
        option_b = str(row[6]).strip() if len(row) > 6 and row[6] else ""
        option_c = str(row[7]).strip() if len(row) > 7 and row[7] else ""
        option_d = str(row[8]).strip() if len(row) > 8 and row[8] else ""
        correct_raw = str(row[9]).strip() if len(row) > 9 and row[9] else ""
        explanation = str(row[10]).strip() if len(row) > 10 and row[10] else None

        if qtype not in {"single", "multiple", "essay"}:
            errors.append(f"Dòng {row_index}: type không hợp lệ ({qtype})")
            continue

        item = QuestionBank(
            teacher_id=teacher_id,
            text=text,
            type=qtype,
            category=category,
            difficulty=difficulty,
            tags=tags,
            explanation=explanation,
        )
        db.session.add(item)
        db.session.flush()

        if qtype != "essay":
            options = [("A", option_a), ("B", option_b), ("C", option_c), ("D", option_d)]
            correct_letters = _extract_correct_letters(correct_raw)
            non_empty_count = 0
            for order, (letter, opt_text) in enumerate(options, start=1):
                if not opt_text:
                    continue
                non_empty_count += 1
                db.session.add(
                    BankOption(
                        question_bank_id=item.id,
                        text=opt_text,
                        is_correct=letter in correct_letters,
                        order_index=order,
                    )
                )
            if non_empty_count < 2:
                errors.append(f"Dòng {row_index}: câu trắc nghiệm cần ít nhất 2 đáp án.")
                db.session.delete(item)
                continue

        imported += 1

    return imported, errors


def _import_from_docx(file_storage, teacher_id: int, default_category: str, default_difficulty: str):
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Thiếu thư viện python-docx. Hãy cài: pip install python-docx") from exc

    document = Document(file_storage)
    lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    imported = 0
    errors = []
    current = None

    def flush_current_question():
        nonlocal imported, current
        if not current:
            return

        question_text = current.get("text", "").strip()
        if not question_text:
            current = None
            return

        qtype = current.get("type", "single")
        item = QuestionBank(
            teacher_id=teacher_id,
            text=question_text,
            type=qtype,
            category=default_category or None,
            difficulty=default_difficulty,
            tags=current.get("tags"),
            explanation=current.get("explanation"),
        )
        db.session.add(item)
        db.session.flush()

        if qtype != "essay":
            options = current.get("options", {})
            correct_letters = _extract_correct_letters(current.get("correct"))
            for idx, letter in enumerate(["A", "B", "C", "D"], start=1):
                if letter in options:
                    db.session.add(
                        BankOption(
                            question_bank_id=item.id,
                            text=options[letter],
                            is_correct=letter in correct_letters,
                            order_index=idx,
                        )
                    )

        imported += 1
        current = None

    for line in lines:
        if line.upper().startswith("Q:"):
            flush_current_question()
            current = {"text": line[2:].strip(), "type": "single", "options": {}}
            continue

        if line.upper().startswith("TYPE:") and current:
            qtype = line.split(":", 1)[1].strip().lower()
            if qtype in {"single", "multiple", "essay"}:
                current["type"] = qtype
            continue

        if line.upper().startswith("ANS:") and current:
            current["correct"] = line.split(":", 1)[1].strip()
            continue

        if line.upper().startswith("EXPL:") and current:
            current["explanation"] = line.split(":", 1)[1].strip()
            continue

        option_match = re.match(r"^([A-Da-d])[\.\):\-]\s*(.+)$", line)
        if option_match and current:
            letter = option_match.group(1).upper()
            current["options"][letter] = option_match.group(2).strip()
            continue

        if current:
            current["text"] = f"{current['text']} {line}".strip()
        else:
            # fallback: mỗi dòng là một câu essay
            item = QuestionBank(
                teacher_id=teacher_id,
                text=line,
                type="essay",
                category=default_category or None,
                difficulty=default_difficulty,
            )
            db.session.add(item)
            imported += 1

    flush_current_question()
    return imported, errors


def _ai_generation_prompt(topic: str, question_count: int, difficulty: str) -> str:
    return f"""
Bạn là NineGPT. Hãy tạo câu hỏi trắc nghiệm cho giáo viên.
Chủ đề: {topic}
Số câu: {question_count}
Độ khó ưu tiên: {difficulty}

Trả về DUY NHẤT một JSON hợp lệ theo schema:
{{
  "questions": [
    {{
      "text": "Nội dung câu hỏi",
      "type": "single|multiple|essay",
      "category": "Tên chủ đề",
      "difficulty": "easy|medium|hard",
      "tags": "tag1,tag2",
      "explanation": "Giải thích ngắn",
      "options": [
        {{"text": "Đáp án 1", "is_correct": true}},
        {{"text": "Đáp án 2", "is_correct": false}},
        {{"text": "Đáp án 3", "is_correct": false}},
        {{"text": "Đáp án 4", "is_correct": false}}
      ]
    }}
  ]
}}

Quy tắc:
- Mỗi câu objective phải có 4 options.
- type=single chỉ 1 đáp án đúng.
- type=multiple có từ 2 đáp án đúng.
- type=essay để options là [].
- Không thêm markdown, không thêm giải thích ngoài JSON.
""".strip()


def _parse_ai_response(raw_text: str) -> list[dict]:
    candidates = [raw_text.strip()]

    fenced = raw_text.strip()
    if fenced.startswith("```"):
        fenced = re.sub(r"^```(?:json)?", "", fenced, flags=re.IGNORECASE).strip()
        fenced = re.sub(r"```$", "", fenced).strip()
        candidates.append(fenced)

    first_obj, last_obj = raw_text.find("{"), raw_text.rfind("}")
    if first_obj != -1 and last_obj != -1 and last_obj > first_obj:
        candidates.append(raw_text[first_obj : last_obj + 1].strip())

    first_arr, last_arr = raw_text.find("["), raw_text.rfind("]")
    if first_arr != -1 and last_arr != -1 and last_arr > first_arr:
        candidates.append(raw_text[first_arr : last_arr + 1].strip())

    for payload in candidates:
        if not payload:
            continue
        try:
            parsed = json.loads(payload)
        except json.JSONDecodeError:
            continue

        if isinstance(parsed, list):
            return [item for item in parsed if isinstance(item, dict)]

        if isinstance(parsed, dict):
            questions = parsed.get("questions") or parsed.get("items")
            if isinstance(questions, list):
                return [item for item in questions if isinstance(item, dict)]

    raise ValueError("AI response không phải JSON hợp lệ theo schema yêu cầu.")


def _call_openai_generate(topic: str, question_count: int, difficulty: str) -> list[dict]:
    api_key = current_app.config.get("OPENAI_API_KEY")
    if not api_key:
        error_msg = "Thiếu OPENAI_API_KEY trong môi trường."
        logger.error(error_msg)
        raise RuntimeError(error_msg)

    model = current_app.config.get("OPENAI_MODEL", "gpt-4o-mini")
    logger.info(f"Calling OpenAI with model={model}, topic={topic}, count={question_count}, difficulty={difficulty}")
    
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": "You generate educational quiz questions in strict JSON."},
            {"role": "user", "content": _ai_generation_prompt(topic, question_count, difficulty)},
        ],
        "temperature": 0.4,
    }
    request_data = json.dumps(payload).encode("utf-8")
    req = urllib_request.Request(
        "https://api.openai.com/v1/chat/completions",
        data=request_data,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )

    try:
        with urllib_request.urlopen(req, timeout=90) as response:
            body = response.read().decode("utf-8")
    except urllib_error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="ignore")
        error_code = exc.code
        
        # Parse error message từ OpenAI response
        error_msg = f"OpenAI API lỗi ({error_code}): {detail[:300]}"
        
        # Xử lý insufficient_quota (429)
        if error_code == 429:
            if "insufficient_quota" in detail.lower():
                error_msg = "❌ Tài khoản OpenAI đã hết quota. Vui lòng nạp credit hoặc kiểm tra billing tại https://platform.openai.com/account/billing/overview"
            else:
                error_msg = "❌ OpenAI API rate limit exceeded. Vui lòng thử lại sau một chút."
        
        logger.error(error_msg)
        raise RuntimeError(error_msg) from exc
    except Exception as exc:
        error_msg = f"Lỗi khi gọi OpenAI: {str(exc)}"
        logger.error(error_msg)
        raise RuntimeError(error_msg) from exc

    response_json = json.loads(body)
    content = response_json.get("choices", [{}])[0].get("message", {}).get("content", "")
    if isinstance(content, list):
        content = "".join(
            part.get("text", "") if isinstance(part, dict) else str(part) for part in content
        )
    logger.info(f"OpenAI returned {len(content)} characters of content")
    return _parse_ai_response(str(content))


def _get_available_gemini_models(api_key: str, api_version: str = "v1") -> list[str]:
    """
    Query Gemini API để lấy danh sách model khả dụng.
    Sử dụng listModels endpoint.
    """
    try:
        url = f"https://generativelanguage.googleapis.com/{api_version}/models?key={urllib_parse.quote(api_key)}"
        req = urllib_request.Request(url, headers={"Content-Type": "application/json"}, method="GET")
        
        with urllib_request.urlopen(req, timeout=10) as response:
            body = response.read().decode("utf-8")
            result = json.loads(body)
        
        # Extract model names
        models = []
        for model_info in result.get("models", []):
            model_name = model_info.get("name", "").split("/")[-1]
            if model_name and "generateContent" in model_info.get("supportedGenerationMethods", []):
                models.append(model_name)
        
        logger.info(f"Available Gemini models on {api_version}: {models}")
        return models
    
    except Exception as exc:
        logger.warning(f"Không thể query danh sách Gemini models: {str(exc)[:200]}")
        return []


def _call_gemini_generate(topic: str, question_count: int, difficulty: str) -> list[dict]:
    """
    Gọi Google Gemini API để sinh câu hỏi.
    - Thử v1 API trước
    - Nếu models không khả dụng trên v1, fallback về v1beta
    - Hỗ trợ fallback sang model khác nếu model hiện tại không hỗ trợ
    """
    api_key = current_app.config.get("GEMINI_API_KEY")
    if not api_key:
        error_msg = "Thiếu GEMINI_API_KEY trong môi trường."
        logger.error(error_msg)
        raise RuntimeError(error_msg)

    model_name = current_app.config.get("GEMINI_MODEL", "gemini-pro")
    api_version = current_app.config.get("GEMINI_API_VERSION", "v1")
    
    # Danh sách API versions để thử (v1 trước, fallback v1beta)
    api_versions_to_try = [api_version]
    if api_version != "v1beta":
        api_versions_to_try.append("v1beta")
    
    for attempt_version in api_versions_to_try:
        logger.info(f"Thử Gemini API version: {attempt_version}")
        
        # Lấy danh sách model khả dụng cho version này
        available_models = _get_available_gemini_models(api_key, attempt_version)
        
        # Nếu không có model nào khả dụng, thử version khác
        if not available_models:
            logger.warning(f"Không có model nào khả dụng trên {attempt_version}")
            continue
        
        # Tạo danh sách model theo độ ưu tiên
        # Ưu tiên: config model → available models → fallback models
        model_fallback_list = []
        
        # Thêm model từ config nếu có
        if model_name in available_models:
            model_fallback_list.append(model_name)
        
        # Thêm available models
        for m in available_models:
            if m not in model_fallback_list:
                model_fallback_list.append(m)
        
        logger.info(f"Model fallback list cho {attempt_version}: {model_fallback_list}")
        
        # Thử từng model
        for attempt_model in model_fallback_list:
            logger.info(f"Gọi Gemini API: model={attempt_model}, version={attempt_version}")
            
            payload = {
                "contents": [
                    {
                        "parts": [
                            {"text": _ai_generation_prompt(topic, question_count, difficulty)}
                        ]
                    }
                ],
                "generationConfig": {
                    "temperature": 0.4,
                },
            }
            
            request_data = json.dumps(payload).encode("utf-8")
            url = f"https://generativelanguage.googleapis.com/{attempt_version}/models/{attempt_model}:generateContent?key={urllib_parse.quote(api_key)}"
            
            req = urllib_request.Request(
                url,
                data=request_data,
                headers={"Content-Type": "application/json"},
                method="POST",
            )

            try:
                with urllib_request.urlopen(req, timeout=90) as response:
                    body = response.read().decode("utf-8")
                    
                response_json = json.loads(body)
                
                # Kiểm tra lỗi trong response
                if "error" in response_json:
                    error_info = response_json["error"]
                    error_code = error_info.get("code")
                    error_msg = error_info.get("message", "Unknown error")
                    
                    logger.warning(f"Gemini model {attempt_model} error: {error_code} - {error_msg}")
                    
                    # Model not found - thử model khác
                    if error_code == 404 and "not found" in error_msg.lower():
                        logger.info(f"Model {attempt_model} không hỗ trợ, thử model khác...")
                        continue
                    
                    # Các lỗi khác
                    raise RuntimeError(f"Gemini API lỗi ({error_code}): {error_msg[:300]}")
                
                # Extract content từ response
                content_text = ""
                candidates = response_json.get("candidates", [])
                if candidates:
                    candidate = candidates[0]
                    if "content" in candidate and "parts" in candidate["content"]:
                        for part in candidate["content"]["parts"]:
                            if "text" in part:
                                content_text += part["text"]
                
                if not content_text:
                    raise RuntimeError("Gemini trả về response trống")
                
                logger.info(f"Gemini model {attempt_model} ({attempt_version}) thành công")
                return _parse_ai_response(content_text)
            
            except urllib_error.HTTPError as exc:
                detail = exc.read().decode("utf-8", errors="ignore")
                error_code = exc.code
                
                logger.warning(f"Gemini model {attempt_model} HTTP error {error_code}")
                
                # Model not found (404) - thử model khác
                if error_code == 404:
                    logger.info(f"Model {attempt_model} không hỗ trợ (404), thử model khác...")
                    continue
                
                # Quota hoặc auth error - không retry
                if error_code in {401, 429}:
                    error_msg = f"Gemini API lỗi ({error_code}): {detail[:300]}"
                    logger.error(error_msg)
                    raise RuntimeError(error_msg) from exc
                
                # Server error - thử model khác
                logger.warning(f"Gemini server error {error_code}, thử model khác...")
                if attempt_model == model_fallback_list[-1]:
                    break  # Chuyển sang version khác
                continue
            
            except Exception as exc:
                logger.warning(f"Lỗi khi gọi Gemini model {attempt_model}: {str(exc)[:200]}")
                if attempt_model == model_fallback_list[-1]:
                    break  # Chuyển sang version khác
                continue
        
        # Nếu version này fail hết, thử version khác
    
    # Nếu hết tất cả API versions và models
    raise RuntimeError(
        f"❌ Không thể kết nối Gemini API. Kiểm tra:\n"
        f"1. GEMINI_API_KEY có hiệu lực không\n"
        f"2. Có model nào khả dụng cho API key này không\n"
        f"3. Thử với OpenAI hoặc chờ và thử lại"
    )



def _persist_ai_questions(
    teacher_id: int,
    generated_questions: list[dict],
    fallback_category: str,
    fallback_difficulty: str,
) -> tuple[int, list[str]]:
    imported = 0
    skipped = []

    for index, item in enumerate(generated_questions, start=1):
        text = str(item.get("text", "")).strip()
        if len(text) < 5:
            skipped.append(f"Câu {index}: thiếu nội dung hợp lệ.")
            continue

        qtype = str(item.get("type", "single")).strip().lower()
        if qtype not in {"single", "multiple", "essay"}:
            qtype = "single"

        category = str(item.get("category") or fallback_category or "").strip() or None
        difficulty = _normalize_difficulty(item.get("difficulty"), default=fallback_difficulty)
        tags = str(item.get("tags", "")).strip() or None
        explanation = str(item.get("explanation", "")).strip() or None

        question = QuestionBank(
            teacher_id=teacher_id,
            text=text,
            type=qtype,
            category=category,
            difficulty=difficulty,
            tags=tags,
            explanation=explanation,
        )
        db.session.add(question)
        db.session.flush()

        if qtype != "essay":
            raw_options = item.get("options")
            if not isinstance(raw_options, list):
                raw_options = []

            normalized_options = []
            for raw in raw_options[:4]:
                if isinstance(raw, dict):
                    opt_text = str(raw.get("text", "")).strip()
                    is_correct = bool(raw.get("is_correct"))
                else:
                    opt_text = str(raw).strip()
                    is_correct = False
                if opt_text:
                    normalized_options.append((opt_text, is_correct))

            if len(normalized_options) < 2:
                db.session.delete(question)
                skipped.append(f"Câu {index}: thiếu đáp án cho câu hỏi objective.")
                continue

            correct_indices = [idx for idx, (_, flag) in enumerate(normalized_options) if flag]
            if not correct_indices:
                first_text = normalized_options[0][0]
                normalized_options[0] = (first_text, True)
                correct_indices = [0]

            if qtype == "single" and len(correct_indices) > 1:
                keep = correct_indices[0]
                normalized_options = [
                    (opt_text, idx == keep) for idx, (opt_text, _) in enumerate(normalized_options)
                ]

            for order, (opt_text, is_correct) in enumerate(normalized_options, start=1):
                db.session.add(
                    BankOption(
                        question_bank_id=question.id,
                        text=opt_text,
                        is_correct=is_correct,
                        order_index=order,
                    )
                )

        imported += 1

    return imported, skipped


def _generate_questions_with_ai(
    provider: str, topic: str, question_count: int, difficulty: str
) -> list[dict]:
    """
    Sinh câu hỏi bằng AI với retry logic và fallback mechanism.
    """
    return _call_ai_with_retry(
        provider=provider,
        topic=topic,
        question_count=question_count,
        difficulty=difficulty,
        attempt=1
    )


@teacher_bp.route("/")
@teacher_required
def dashboard():
    my_quiz_ids = [quiz.id for quiz in Quiz.query.filter_by(teacher_id=current_user.id).all()]
    total_questions = QuestionBank.query.filter_by(teacher_id=current_user.id).count()
    total_quizzes = len(my_quiz_ids)

    if my_quiz_ids:
        recent_results = (
            Submission.query.filter(
                Submission.quiz_id.in_(my_quiz_ids),
                Submission.status.in_(["submitted", "graded"]),
            )
            .order_by(Submission.submitted_at.desc())
            .limit(10)
            .all()
        )
    else:
        recent_results = []

    recent_quizzes = (
        Quiz.query.filter_by(teacher_id=current_user.id)
        .order_by(Quiz.created_at.desc())
        .limit(5)
        .all()
    )

    avg_score = (
        round(sum(row.score_percent for row in recent_results) / len(recent_results), 1)
        if recent_results
        else 0
    )
    total_students = len(
        {
            row.student_id or row.display_student_email.lower()
            for row in recent_results
            if row.student_id or row.display_student_email
        }
    )

    return render_template(
        "teacher/dashboard_v2.html",
        total_questions=total_questions,
        total_quizzes=total_quizzes,
        total_students=total_students,
        total_results=len(recent_results),
        recent_quizzes=recent_quizzes,
        recent_results=recent_results,
        avg_score=avg_score,
    )


# ==================== CLASSROOM MANAGEMENT ====================
@teacher_bp.route("/classes")
@teacher_required
def classes_list():
    classes = Classroom.query.filter_by(teacher_id=current_user.id).order_by(Classroom.created_at.desc()).all()
    return render_template("teacher/classes/list.html", classes=classes)


@teacher_bp.route("/classes/create", methods=["GET", "POST"])
@teacher_required
def classes_create():
    form = ClassroomForm()
    if form.validate_on_submit():
        classroom = Classroom(
            name=form.name.data,
            description=form.description.data,
            teacher_id=current_user.id,
            join_code=Classroom.generate_join_code(),
        )
        db.session.add(classroom)
        db.session.commit()
        flash("Tạo lớp học thành công.", "success")
        return redirect(url_for("teacher.classes_list"))
    return render_template("teacher/classes/form.html", form=form, title="Tạo lớp học")


@teacher_bp.route("/classes/<int:class_id>/edit", methods=["GET", "POST"])
@teacher_required
def classes_edit(class_id):
    classroom = Classroom.query.get_or_404(class_id)
    if classroom.teacher_id != current_user.id:
        flash("Bạn không có quyền sửa lớp này.", "danger")
        return redirect(url_for("teacher.classes_list"))

    form = ClassroomForm(obj=classroom)
    if form.validate_on_submit():
        classroom.name = form.name.data
        classroom.description = form.description.data
        db.session.commit()
        flash("Cập nhật lớp học thành công.", "success")
        return redirect(url_for("teacher.classes_list"))
    return render_template("teacher/classes/form.html", form=form, title="Sửa lớp học")


@teacher_bp.route("/classes/<int:class_id>/delete", methods=["POST"])
@teacher_required
def classes_delete(class_id):
    classroom = Classroom.query.get_or_404(class_id)
    if classroom.teacher_id != current_user.id:
        flash("Bạn không có quyền xóa lớp này.", "danger")
        return redirect(url_for("teacher.classes_list"))

    db.session.delete(classroom)
    db.session.commit()
    flash("Đã xóa lớp học.", "success")
    return redirect(url_for("teacher.classes_list"))


@teacher_bp.route("/classes/<int:class_id>/members")
@teacher_required
def class_members(class_id):
    classroom = Classroom.query.get_or_404(class_id)
    if classroom.teacher_id != current_user.id:
        flash("Bạn không có quyền xem lớp này.", "danger")
        return redirect(url_for("teacher.classes_list"))

    pending_members = (
        UserClass.query.filter_by(class_id=classroom.id, status="pending")
        .order_by(UserClass.enrolled_at.asc())
        .all()
    )
    approved_members = (
        UserClass.query.filter_by(class_id=classroom.id, status="approved")
        .order_by(UserClass.enrolled_at.asc())
        .all()
    )
    documents = (
        ClassDocument.query.filter_by(class_id=classroom.id)
        .order_by(ClassDocument.uploaded_at.desc())
        .all()
    )
    upload_form = DocumentUploadForm()
    join_link = url_for("quiz.join_class_by_link", join_code=classroom.join_code, _external=True)
    class_qr_url = f"https://api.qrserver.com/v1/create-qr-code/?size=220x220&data={quote(join_link)}"
    return render_template(
        "teacher/classes/members.html",
        classroom=classroom,
        pending_members=pending_members,
        approved_members=approved_members,
        documents=documents,
        upload_form=upload_form,
        join_link=join_link,
        class_qr_url=class_qr_url,
    )


@teacher_bp.route("/classes/<int:class_id>/members/<int:membership_id>/approve", methods=["POST"])
@teacher_required
def class_member_approve(class_id, membership_id):
    classroom = Classroom.query.get_or_404(class_id)
    membership = UserClass.query.get_or_404(membership_id)

    if classroom.teacher_id != current_user.id or membership.class_id != classroom.id:
        flash("Không thể duyệt thành viên.", "danger")
        return redirect(url_for("teacher.class_members", class_id=class_id))

    membership.status = "approved"
    db.session.commit()
    flash("Đã duyệt học sinh vào lớp.", "success")
    return redirect(url_for("teacher.class_members", class_id=class_id))


@teacher_bp.route("/classes/<int:class_id>/members/<int:membership_id>/remove", methods=["POST"])
@teacher_required
def class_member_remove(class_id, membership_id):
    classroom = Classroom.query.get_or_404(class_id)
    membership = UserClass.query.get_or_404(membership_id)

    if classroom.teacher_id != current_user.id or membership.class_id != classroom.id:
        flash("Không thể xóa thành viên.", "danger")
        return redirect(url_for("teacher.class_members", class_id=class_id))

    db.session.delete(membership)
    db.session.commit()
    flash("Đã xóa học sinh khỏi lớp.", "success")
    return redirect(url_for("teacher.class_members", class_id=class_id))


@teacher_bp.route("/classes/<int:class_id>/documents/upload", methods=["POST"])
@teacher_required
def class_document_upload(class_id):
    classroom = Classroom.query.get_or_404(class_id)
    if classroom.teacher_id != current_user.id:
        flash("Bạn không có quyền upload tài liệu cho lớp này.", "danger")
        return redirect(url_for("teacher.classes_list"))

    form = DocumentUploadForm()
    if not form.validate_on_submit():
        flash("Thông tin tài liệu chưa hợp lệ.", "danger")
        return redirect(url_for("teacher.class_members", class_id=class_id))

    file = form.file.data
    filename = secure_filename(file.filename or "")
    if not filename:
        flash("Tên file không hợp lệ.", "danger")
        return redirect(url_for("teacher.class_members", class_id=class_id))

    upload_root = os.path.join(current_app.root_path, "static", "uploads", "documents", str(current_user.id))
    os.makedirs(upload_root, exist_ok=True)
    file_path = os.path.join(upload_root, filename)
    file.save(file_path)

    file_url = url_for(
        "static",
        filename=f"uploads/documents/{current_user.id}/{filename}",
        _external=False,
    )
    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    document = ClassDocument(
        class_id=class_id,
        teacher_id=current_user.id,
        title=form.title.data,
        file_url=file_url,
        file_type=extension,
    )
    db.session.add(document)
    db.session.commit()
    flash("Đã upload tài liệu lớp học.", "success")
    return redirect(url_for("teacher.class_members", class_id=class_id))


# ==================== QUESTION BANK ====================
@teacher_bp.route("/questions")
@teacher_required
def questions_list():
    search = request.args.get("search", "").strip()
    category = request.args.get("category", "").strip()
    difficulty = request.args.get("difficulty", "").strip()

    query = QuestionBank.query.filter_by(teacher_id=current_user.id)
    if search:
        query = query.filter(QuestionBank.text.ilike(f"%{search}%"))
    if category:
        query = query.filter(QuestionBank.category.ilike(f"%{category}%"))
    if difficulty in {"easy", "medium", "hard"}:
        query = query.filter_by(difficulty=difficulty)

    questions = query.order_by(QuestionBank.created_at.desc()).all()
    categories = (
        db.session.query(QuestionBank.category)
        .filter(QuestionBank.teacher_id == current_user.id, QuestionBank.category.isnot(None))
        .distinct()
        .all()
    )
    categories = [c[0] for c in categories if c[0]]

    return render_template(
        "teacher/questions/list.html",
        questions=questions,
        categories=categories,
        current_search=search,
        current_category=category,
        current_difficulty=difficulty,
    )


@teacher_bp.route("/questions/create", methods=["GET", "POST"])
@teacher_required
def questions_create():
    form = QuestionBankForm()
    if form.validate_on_submit():
        question = QuestionBank(
            teacher_id=current_user.id,
            text=form.text.data,
            type=form.type.data,
            category=form.category.data or None,
            difficulty=form.difficulty.data,
            tags=form.tags.data or None,
            explanation=form.explanation.data or None,
        )
        db.session.add(question)
        db.session.flush()
        _sync_bank_options(question, form)
        db.session.commit()
        flash("Đã tạo câu hỏi trong ngân hàng.", "success")
        return redirect(url_for("teacher.questions_list"))

    return render_template("teacher/questions/create.html", form=form)


@teacher_bp.route("/questions/<int:id>/edit", methods=["GET", "POST"])
@teacher_required
def questions_edit(id):
    question = QuestionBank.query.get_or_404(id)
    if question.teacher_id != current_user.id:
        flash("Bạn không có quyền sửa câu hỏi này.", "danger")
        return redirect(url_for("teacher.questions_list"))

    form = QuestionBankForm(obj=question)
    existing = {opt.order_index: opt for opt in question.options.all()}
    if request.method == "GET":
        form.option_a.data = existing.get(1).text if existing.get(1) else ""
        form.option_b.data = existing.get(2).text if existing.get(2) else ""
        form.option_c.data = existing.get(3).text if existing.get(3) else ""
        form.option_d.data = existing.get(4).text if existing.get(4) else ""
        form.correct_answers.data = ",".join([opt.label for opt in question.options.filter_by(is_correct=True)])

    if form.validate_on_submit():
        question.text = form.text.data
        question.type = form.type.data
        question.category = form.category.data or None
        question.difficulty = form.difficulty.data
        question.tags = form.tags.data or None
        question.explanation = form.explanation.data or None
        _sync_bank_options(question, form)
        db.session.commit()
        flash("Cập nhật câu hỏi thành công.", "success")
        return redirect(url_for("teacher.questions_list"))

    return render_template("teacher/questions/edit.html", form=form, question=question)


@teacher_bp.route("/questions/<int:id>/delete", methods=["POST"])
@teacher_required
def questions_delete(id):
    question = QuestionBank.query.get_or_404(id)
    if question.teacher_id != current_user.id:
        flash("Bạn không có quyền xóa câu hỏi này.", "danger")
        return redirect(url_for("teacher.questions_list"))

    db.session.delete(question)
    db.session.commit()
    flash("Đã xóa câu hỏi khỏi ngân hàng.", "success")
    return redirect(url_for("teacher.questions_list"))


@teacher_bp.route("/questions/import", methods=["GET", "POST"])
@teacher_required
def questions_import():
    form = ImportBankForm()
    if form.validate_on_submit():
        file = form.file.data
        filename = (file.filename or "").lower()
        default_category = form.default_category.data.strip() if form.default_category.data else ""
        default_difficulty = _normalize_difficulty(form.default_difficulty.data)

        try:
            if filename.endswith((".xlsx", ".xls")):
                imported, errors = _import_from_excel(
                    file, current_user.id, default_category, default_difficulty
                )
            elif filename.endswith(".docx"):
                imported, errors = _import_from_docx(
                    file, current_user.id, default_category, default_difficulty
                )
            else:
                flash("File không hợp lệ. Chỉ hỗ trợ .xlsx/.xls/.docx", "danger")
                return render_template("teacher/questions/import.html", form=form)

            db.session.commit()
            if imported:
                flash(f"Import thành công {imported} câu hỏi.", "success")
            if errors:
                flash("Một số dòng bị bỏ qua: " + "; ".join(errors[:5]), "warning")
            return redirect(url_for("teacher.questions_list"))
        except Exception as exc:
            db.session.rollback()
            flash(f"Lỗi import: {exc}", "danger")

    return render_template("teacher/questions/import.html", form=form)


@teacher_bp.route("/api/ai-provider-status", methods=["GET"])
@teacher_required
def api_ai_provider_status():
    """
    API endpoint để kiểm tra trạng thái của các AI providers.
    Trả về JSON với status của OpenAI và Gemini.
    """
    status = _get_ai_provider_status()
    return jsonify(status)


@teacher_bp.route("/questions/ai-generate", methods=["GET", "POST"])
@teacher_required
def questions_ai_generate():
    form = AIGenerateQuestionsForm()
    
    # Xử lý GET request - hiển thị form
    if request.method == "GET":
        return render_template("teacher/questions/ai_generate.html", form=form)
    
    # Xử lý POST request
    if request.is_json:
        # JSON request từ AJAX
        data = request.get_json() or {}
        provider = data.get("provider", "openai").strip()
        topic = data.get("topic", "").strip()
        category = data.get("category", "").strip() or topic
        difficulty = _normalize_difficulty(data.get("difficulty", "medium"))
        question_count = data.get("question_count", 5)
        
        # Validation
        if not topic:
            return jsonify({"success": False, "error": "Vui lòng nhập chủ đề"}), 400
        
        if len(topic) < 3:
            return jsonify({"success": False, "error": "Chủ đề phải có ít nhất 3 ký tự"}), 400
        
        if not isinstance(question_count, int) or question_count < 1 or question_count > 10:
            return jsonify({"success": False, "error": "Số câu phải từ 1-10"}), 400
        
        if provider not in {"openai", "gemini"}:
            return jsonify({"success": False, "error": "AI provider không hợp lệ"}), 400
    
    else:
        # Form data request
        if not form.validate_on_submit():
            # Return form errors
            errors = {field: [error for error in form[field].errors] for field in form.errors}
            return jsonify({"success": False, "error": "Form validation failed", "errors": errors}), 400
        
        provider = form.provider.data
        topic = form.topic.data.strip()
        category = form.category.data.strip() if form.category.data else topic
        difficulty = form.difficulty.data
        question_count = form.question_count.data
    
    try:
        logger.info(f"Generating AI questions: provider={provider}, topic={topic}, count={question_count}, difficulty={difficulty}")
        
        # Gọi AI để sinh câu hỏi
        generated_questions = _generate_questions_with_ai(
            provider=provider,
            topic=topic,
            question_count=question_count,
            difficulty=difficulty,
        )
        
        if not generated_questions:
            return jsonify({"success": False, "error": "AI không trả về câu hỏi nào. Thử lại với chủ đề khác."}), 400
        
        logger.info(f"Generated {len(generated_questions)} questions for user {current_user.id}")
        
        # Trả về JSON để frontend hiển thị
        return jsonify({
            "success": True,
            "questions": generated_questions,
            "category": category,
            "difficulty": difficulty,
            "topic": topic
        })
        
    except Exception as exc:
        db.session.rollback()
        error_msg = str(exc)
        logger.error(f"Error generating AI questions: {error_msg}")
        
        # Trích xuất thông báo lỗi hữu ích từ exception
        if "OPENAI_API_KEY" in error_msg or "GEMINI_API_KEY" in error_msg:
            error_display = "❌ Chưa cấu hình API key cho AI provider này. Hãy liên hệ quản trị viên."
        elif "quota" in error_msg.lower():
            error_display = "❌ Quota API đã hết. Vui lòng thử lại sau hoặc chuyển sang provider khác."
        elif "timeout" in error_msg.lower():
            error_display = "❌ Kết nối AI hết thời gian chờ. Vui lòng thử lại."
        else:
            error_display = f"❌ Lỗi AI: {error_msg[:200]}"
        
        return jsonify({"success": False, "error": error_display}), 400


@teacher_bp.route("/questions/save-ai-batch", methods=["POST"])
@teacher_required
def save_ai_questions_batch():
    """
    Lưu hàng loạt câu hỏi từ AI vào database.
    Request body: {questions: [...], category, difficulty}
    Response: {success: bool, saved_count: int, error: str}
    """
    try:
        data = request.get_json()
        questions = data.get("questions", [])
        fallback_category = data.get("category")
        fallback_difficulty = data.get("difficulty", "medium")
        
        if not questions:
            return jsonify({"success": False, "error": "Không có câu hỏi nào"})
        
        imported, skipped = _persist_ai_questions(
            teacher_id=current_user.id,
            generated_questions=questions,
            fallback_category=fallback_category,
            fallback_difficulty=fallback_difficulty,
        )
        db.session.commit()
        
        if imported > 0:
            logger.info(f"Saved {imported} AI questions for user {current_user.id}")
            return jsonify({"success": True, "saved_count": imported, "skipped": skipped})
        else:
            return jsonify({"success": False, "error": "Tất cả câu hỏi đều bị bỏ qua"})
            
    except Exception as exc:
        db.session.rollback()
        logger.error(f"Error saving AI questions: {exc}")
        return jsonify({"success": False, "error": str(exc)})


@teacher_bp.route("/questions/create-quiz-from-ai", methods=["POST"])
@teacher_required
def create_quiz_from_ai():
    """
    Tạo quiz mới và thêm câu hỏi từ AI vào đó.
    Request body: {quiz_name, questions: [...]}
    Response: {success: bool, redirect_url: str, error: str}
    """
    try:
        data = request.get_json()
        quiz_name = data.get("quiz_name", "Bài kiểm tra từ AI").strip()
        questions = data.get("questions", [])
        
        if not questions:
            return jsonify({"success": False, "error": "Không có câu hỏi nào"})
        
        # Tạo quiz mới
        new_quiz = Quiz(
            title=quiz_name,
            teacher_id=current_user.id,
            description=f"Tạo tự động bởi NineGPT - {len(questions)} câu hỏi",
            time_limit_minutes=60,  # Cột mới
            is_shuffled=False,
            shuffle_options=False,
            is_published=False,
            allow_retake=False,
        )
        db.session.add(new_quiz)
        db.session.flush()  # Lấy ID của quiz
        
        # Thêm câu hỏi vào quiz
        for order_index, q_data in enumerate(questions, start=1):
            # Tạo câu hỏi trong database
            question = Question(
                quiz_id=new_quiz.id,
                text=q_data.get("text", ""),
                type=q_data.get("type", "single"),
                points=1,
                order_index=order_index,
                topic=q_data.get("category"),
                difficulty=q_data.get("difficulty", "medium"),
                created_by=current_user.id,
                explanation=q_data.get("explanation"),
            )
            db.session.add(question)
            db.session.flush()
            
            # Thêm đáp án
            if q_data.get("type") != "essay":
                for opt_idx, opt_data in enumerate(q_data.get("options", []), start=1):
                    option = Option(
                        question_id=question.id,
                        text=opt_data.get("text", ""),
                        is_correct=opt_data.get("is_correct", False),
                        order_index=opt_idx,
                    )
                    db.session.add(option)
        
        db.session.commit()
        logger.info(f"Created quiz {new_quiz.id} from AI with {len(questions)} questions")
        
        return jsonify({
            "success": True,
            "redirect_url": url_for("teacher.quizzes_edit", id=new_quiz.id)
        })
        
    except Exception as exc:
        db.session.rollback()
        logger.error(f"Error creating quiz from AI: {exc}")
        return jsonify({"success": False, "error": str(exc)})


@teacher_bp.route("/questions/export")
@teacher_required
def questions_export():
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "QuestionBank"
    sheet.append(
        [
            "text",
            "type",
            "category",
            "difficulty",
            "tags",
            "option_a",
            "option_b",
            "option_c",
            "option_d",
            "correct_answers",
            "explanation",
        ]
    )

    rows = QuestionBank.query.filter_by(teacher_id=current_user.id).order_by(QuestionBank.id.asc()).all()
    for item in rows:
        options = {opt.order_index: opt for opt in item.options.all()}
        correct = ",".join([opt.label for opt in item.options.filter_by(is_correct=True).all()])
        sheet.append(
            [
                item.text,
                item.type,
                item.category or "",
                item.difficulty,
                item.tags or "",
                options.get(1).text if options.get(1) else "",
                options.get(2).text if options.get(2) else "",
                options.get(3).text if options.get(3) else "",
                options.get(4).text if options.get(4) else "",
                correct,
                item.explanation or "",
            ]
        )

    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)
    return send_file(
        output,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=f"question_bank_{current_user.id}.xlsx",
    )


@teacher_bp.route("/questions/template")
@teacher_required
def questions_template():
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Template"
    sheet.append(
        [
            "text",
            "type",
            "category",
            "difficulty",
            "tags",
            "option_a",
            "option_b",
            "option_c",
            "option_d",
            "correct_answers",
            "explanation",
        ]
    )
    sheet.append(
        [
            "Python là ngôn ngữ gì?",
            "single",
            "Tin học",
            "easy",
            "python,co-ban",
            "Ngôn ngữ biên dịch",
            "Ngôn ngữ thông dịch",
            "Ngôn ngữ máy",
            "Ngôn ngữ đánh dấu",
            "B",
            "Python là ngôn ngữ thông dịch bậc cao.",
        ]
    )

    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)
    return send_file(
        output,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name="question_bank_template.xlsx",
    )


# ==================== QUIZ MANAGEMENT ====================
@teacher_bp.route("/quizzes")
@teacher_required
def quizzes_list():
    quizzes = (
        Quiz.query.filter_by(teacher_id=current_user.id)
        .order_by(Quiz.created_at.desc())
        .all()
    )
    return render_template("teacher/quizzes/list.html", quizzes=quizzes)


@teacher_bp.route("/quizzes/create", methods=["GET", "POST"])
@teacher_required
def quizzes_create():
    form = QuizBuilderForm()
    class_choices = [(0, "Không gán lớp")] + [
        (classroom.id, f"{classroom.name} ({classroom.join_code})")
        for classroom in Classroom.query.filter_by(teacher_id=current_user.id).all()
    ]
    form.class_id.choices = class_choices

    available_bank_questions = QuestionBank.query.filter_by(teacher_id=current_user.id).all()

    if form.validate_on_submit():
        scheduled_at = _parse_datetime(form.scheduled_at.data)
        expires_at = _parse_datetime(form.expires_at.data)

        quiz = Quiz(
            title=form.title.data,
            description=form.description.data or None,
            time_limit_minutes=form.time_limit_minutes.data,
            is_shuffled=form.is_shuffled.data,
            shuffle_options=form.shuffle_options.data,
            is_published=form.is_published.data,
            allow_retake=form.allow_retake.data,
            teacher_id=current_user.id,
            class_id=form.class_id.data if form.class_id.data else None,
            scheduled_at=scheduled_at,
            expires_at=expires_at,
            brand_logo_url=form.brand_logo_url.data or None,
            brand_color=form.brand_color.data or None,
        )
        db.session.add(quiz)
        db.session.flush()

        selected_bank_questions = []
        if form.source_mode.data == "random":
            random_query = QuestionBank.query.filter_by(teacher_id=current_user.id)
            if form.random_category.data:
                random_query = random_query.filter(
                    QuestionBank.category.ilike(f"%{form.random_category.data.strip()}%")
                )
            if form.random_difficulty.data:
                random_query = random_query.filter_by(difficulty=form.random_difficulty.data)
            candidates = random_query.all()
            requested_count = form.random_count.data or 10
            selected_bank_questions = (
                random.sample(candidates, requested_count)
                if len(candidates) >= requested_count
                else candidates
            )
        else:
            selected_ids = [
                int(raw_id)
                for raw_id in request.form.getlist("bank_question_ids")
                if raw_id.isdigit()
            ]
            if selected_ids:
                selected_bank_questions = (
                    QuestionBank.query.filter(
                        QuestionBank.teacher_id == current_user.id,
                        QuestionBank.id.in_(selected_ids),
                    )
                    .order_by(QuestionBank.id.asc())
                    .all()
                )

        for idx, bank_item in enumerate(selected_bank_questions, start=1):
            db.session.add(bank_item.clone_to_quiz(quiz.id, order_index=idx, created_by=current_user.id))

        db.session.commit()
        flash(f"Tạo đề thi thành công với {len(selected_bank_questions)} câu.", "success")
        return redirect(url_for("teacher.quizzes_list"))

    return render_template(
        "teacher/quizzes/create.html",
        form=form,
        bank_questions=available_bank_questions,
    )


@teacher_bp.route("/quizzes/<int:id>/edit", methods=["GET", "POST"])
@teacher_required
def quizzes_edit(id):
    quiz = Quiz.query.get_or_404(id)
    if quiz.teacher_id != current_user.id:
        flash("Bạn không có quyền sửa đề thi này.", "danger")
        return redirect(url_for("teacher.quizzes_list"))

    form = QuizBuilderForm(obj=quiz)
    class_choices = [(0, "Không gán lớp")] + [
        (classroom.id, f"{classroom.name} ({classroom.join_code})")
        for classroom in Classroom.query.filter_by(teacher_id=current_user.id).all()
    ]
    form.class_id.choices = class_choices

    if request.method == "GET":
        form.class_id.data = quiz.class_id or 0
        form.scheduled_at.data = quiz.scheduled_at.strftime("%Y-%m-%d %H:%M") if quiz.scheduled_at else ""
        form.expires_at.data = quiz.expires_at.strftime("%Y-%m-%d %H:%M") if quiz.expires_at else ""

    if form.validate_on_submit():
        quiz.title = form.title.data
        quiz.description = form.description.data or None
        quiz.time_limit_minutes = form.time_limit_minutes.data
        quiz.is_shuffled = form.is_shuffled.data
        quiz.shuffle_options = form.shuffle_options.data
        quiz.is_published = form.is_published.data
        quiz.allow_retake = form.allow_retake.data
        quiz.class_id = form.class_id.data if form.class_id.data else None
        quiz.scheduled_at = _parse_datetime(form.scheduled_at.data)
        quiz.expires_at = _parse_datetime(form.expires_at.data)
        quiz.brand_logo_url = form.brand_logo_url.data or None
        quiz.brand_color = form.brand_color.data or None

        add_ids = [int(raw_id) for raw_id in request.form.getlist("bank_question_ids") if raw_id.isdigit()]
        if add_ids:
            existing_count = quiz.questions.count()
            additions = (
                QuestionBank.query.filter(
                    QuestionBank.teacher_id == current_user.id,
                    QuestionBank.id.in_(add_ids),
                )
                .order_by(QuestionBank.id.asc())
                .all()
            )
            for offset, bank_item in enumerate(additions, start=1):
                db.session.add(
                    bank_item.clone_to_quiz(
                        quiz.id,
                        order_index=existing_count + offset,
                        created_by=current_user.id,
                    )
                )

        db.session.commit()
        flash("Cập nhật đề thi thành công.", "success")
        return redirect(url_for("teacher.quizzes_list"))

    bank_questions = QuestionBank.query.filter_by(teacher_id=current_user.id).all()
    quiz_questions = quiz.get_ordered_questions()
    return render_template(
        "teacher/quizzes/edit.html",
        form=form,
        quiz=quiz,
        bank_questions=bank_questions,
        quiz_questions=quiz_questions,
    )


@teacher_bp.route("/quizzes/<int:id>/questions/<int:question_id>/media", methods=["POST"])
@teacher_required
def quiz_question_media_update(id, question_id):
    quiz = Quiz.query.get_or_404(id)
    if quiz.teacher_id != current_user.id:
        flash("Bạn không có quyền cập nhật đề này.", "danger")
        return redirect(url_for("teacher.quizzes_list"))

    question = Question.query.get_or_404(question_id)
    if question.quiz_id != quiz.id:
        flash("Câu hỏi không thuộc đề thi này.", "danger")
        return redirect(url_for("teacher.quizzes_edit", id=id))

    image_url = (request.form.get("image_url") or "").strip() or None
    youtube_url = (request.form.get("youtube_url") or "").strip() or None

    if youtube_url and "youtube.com" not in youtube_url and "youtu.be" not in youtube_url:
        flash("Link YouTube không hợp lệ.", "danger")
        return redirect(url_for("teacher.quizzes_edit", id=id))

    question.image_url = image_url
    question.youtube_url = youtube_url
    db.session.commit()
    flash("Đã cập nhật media cho câu hỏi.", "success")
    return redirect(url_for("teacher.quizzes_edit", id=id))


@teacher_bp.route("/quizzes/<int:id>/delete", methods=["POST"])
@teacher_required
def quizzes_delete(id):
    quiz = Quiz.query.get_or_404(id)
    if quiz.teacher_id != current_user.id:
        flash("Bạn không có quyền xóa đề này.", "danger")
        return redirect(url_for("teacher.quizzes_list"))
    db.session.delete(quiz)
    db.session.commit()
    flash("Đã xóa đề thi.", "success")
    return redirect(url_for("teacher.quizzes_list"))


@teacher_bp.route("/quizzes/<int:id>/regenerate", methods=["POST"])
@teacher_required
def quizzes_regenerate(id):
    quiz = Quiz.query.get_or_404(id)
    if quiz.teacher_id != current_user.id:
        flash("Bạn không có quyền tạo lại đề này.", "danger")
        return redirect(url_for("teacher.quizzes_list"))

    random_count = request.form.get("random_count", type=int) or max(quiz.questions.count(), 5)
    candidates = QuestionBank.query.filter_by(teacher_id=current_user.id).all()
    selected = random.sample(candidates, random_count) if len(candidates) >= random_count else candidates

    quiz.questions.delete()
    for idx, bank_item in enumerate(selected, start=1):
        db.session.add(bank_item.clone_to_quiz(quiz.id, order_index=idx, created_by=current_user.id))
    db.session.commit()

    flash(f"Đã tạo lại đề thi với {len(selected)} câu hỏi.", "success")
    return redirect(url_for("teacher.quizzes_edit", id=id))


@teacher_bp.route("/quizzes/<int:id>/export")
@teacher_required
def quizzes_export(id):
    quiz = Quiz.query.get_or_404(id)
    if quiz.teacher_id != current_user.id:
        flash("Bạn không có quyền export đề này.", "danger")
        return redirect(url_for("teacher.quizzes_list"))

    workbook = Workbook()
    info_sheet = workbook.active
    info_sheet.title = "Quiz"
    info_sheet.append(["Title", quiz.title])
    info_sheet.append(["Description", quiz.description or ""])
    info_sheet.append(["Class", quiz.classroom.name if quiz.classroom else ""])
    info_sheet.append(["TimeLimit", quiz.time_limit_minutes])
    info_sheet.append(["ShuffleQuestions", "Yes" if quiz.is_shuffled else "No"])
    info_sheet.append(["ShuffleOptions", "Yes" if quiz.shuffle_options else "No"])
    info_sheet.append(["Published", "Yes" if quiz.is_published else "No"])

    question_sheet = workbook.create_sheet("Questions")
    question_sheet.append(["Order", "Type", "Question", "CorrectOptions", "Options", "Explanation"])
    for question in quiz.get_ordered_questions():
        options = question.options.order_by("order_index").all()
        options_text = " | ".join([f"{opt.label}. {opt.text}" for opt in options])
        correct_letters = ",".join([opt.label for opt in question.correct_options])
        question_sheet.append(
            [
                question.order_index,
                question.type,
                question.text,
                correct_letters,
                options_text,
                question.explanation or "",
            ]
        )

    result_sheet = workbook.create_sheet("Submissions")
    result_sheet.append(["Student", "Email", "Score", "MaxScore", "Percent", "Status", "SubmittedAt"])
    for sub in quiz.submissions.order_by(Submission.submitted_at.desc()).all():
        result_sheet.append(
            [
                sub.display_student_name,
                sub.display_student_email,
                sub.total_score,
                sub.max_score,
                sub.score_percent,
                sub.status,
                sub.submitted_at.strftime("%d/%m/%Y %H:%M") if sub.submitted_at else "",
            ]
        )

    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)
    return send_file(
        output,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=f"quiz_{quiz.id}.xlsx",
    )


@teacher_bp.route("/quizzes/<int:id>/share")
@teacher_required
def quizzes_share(id):
    quiz = Quiz.query.get_or_404(id)
    if quiz.teacher_id != current_user.id:
        flash("Bạn không có quyền chia sẻ đề này.", "danger")
        return redirect(url_for("teacher.quizzes_list"))

    token = quiz.ensure_share_token()
    db.session.commit()

    share_url = url_for("quiz.shared_quiz_entry", token=token, _external=True)
    qr_url = f"https://api.qrserver.com/v1/create-qr-code/?size=260x260&data={quote(share_url)}"
    return render_template("teacher/quizzes/share.html", quiz=quiz, share_url=share_url, qr_url=qr_url)


# ==================== RESULTS & MANUAL GRADING ====================
@teacher_bp.route("/results")
@teacher_required
def results():
    my_quiz_ids = [quiz.id for quiz in Quiz.query.filter_by(teacher_id=current_user.id).all()]
    if not my_quiz_ids:
        submissions = []
    else:
        submissions = (
            Submission.query.filter(Submission.quiz_id.in_(my_quiz_ids))
            .order_by(Submission.submitted_at.desc())
            .all()
        )

    return render_template("teacher/results/list.html", results=submissions)


@teacher_bp.route("/results/<int:result_id>")
@teacher_required
def result_detail(result_id):
    submission = Submission.query.get_or_404(result_id)
    if submission.quiz.teacher_id != current_user.id:
        flash("Bạn không có quyền xem kết quả này.", "danger")
        return redirect(url_for("teacher.results"))

    grading_forms = {answer.id: EssayGradingForm() for answer in submission.answers.all()}
    return render_template(
        "teacher/results/detail.html",
        result=submission,
        user_answers=submission.answers.order_by(Answer.id.asc()).all(),
        grading_forms=grading_forms,
    )


@teacher_bp.route("/results/<int:result_id>/grade/<int:answer_id>", methods=["POST"])
@teacher_required
def grade_essay(result_id, answer_id):
    submission = Submission.query.get_or_404(result_id)
    if submission.quiz.teacher_id != current_user.id:
        flash("Bạn không có quyền chấm bài này.", "danger")
        return redirect(url_for("teacher.results"))

    answer = Answer.query.get_or_404(answer_id)
    if answer.submission_id != submission.id or answer.question.type != "essay":
        flash("Không thể chấm câu trả lời này.", "danger")
        return redirect(url_for("teacher.result_detail", result_id=result_id))

    form = EssayGradingForm()
    if form.validate_on_submit():
        answer.apply_manual_grading(score=form.score_awarded.data, feedback=form.feedback.data)
        answer.is_correct = answer.score_awarded >= max(answer.question.points, 1)

        submission.status = "graded"
        submission.recalculate_scores()
        if not submission.submitted_at:
            submission.submitted_at = datetime.utcnow()
        db.session.commit()
        flash("Đã lưu chấm điểm tự luận.", "success")
    else:
        flash("Thông tin chấm điểm chưa hợp lệ.", "danger")

    return redirect(url_for("teacher.result_detail", result_id=result_id))

