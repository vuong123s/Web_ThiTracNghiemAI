from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

content_vn = '''Tóm tắt  –  Trong thời đại kỷ nguyên số, việc ứng dụng trí tuệ nhân tạo vào giáo dục đang mở ra những cơ hội mới để cải thiện chất lượng giảng dạy và quản lý lớp học. Với mục tiêu nâng cao hiệu quả trong công tác quản lý lớp học, nghiên cứu này của chúng tôi đề xuất để xây dựng một ứng dụng hỗ trợ quản lý lớp học, nhằm tự động hóa một số quy trình trong quản lý lớp học và cải thiện hiệu suất quản lý. Ứng dụng bao gồm chức năng chính là điểm danh bằng phương pháp nhận diện khuôn mặt sử dụng mô hình MTCNN kết hợp MobileFaceNet với độ chính xác 99.55% trên tập dữ liệu LFW. Mô hình được được chỉnh sửa ngưỡng threshold để phù hợp với điều kiện môi trường trong nhận diện khuôn mặt thời gian thực với thời gian xử lý trung bình 0.01 giây khi sử dụng CPU Core I7 10750H, giúp giảm thiểu sự sai sót và thời gian so với các phương pháp điểm danh truyền thống còn tồn đọng. Đồng thời ứng dụng cũng sử dụng mô hình YOLO v11 để nhận diện người  giúp giảng viên có theo dõi và xác định số lượng sinh viên đang có trong phòng học thời điểm hiện tại thông qua các camera giám sát, đạt độ chính xác khoảng 85% trên thang điểm mAP@50-95. Cung cấp ứng dụng web giúp giảng viên trực quan hóa cáo báo cáo, thống kê về tình hình điểm danh của sinh viên, cho phép giảng viên giám sát lớp học, giám sát quy trình điểm danh theo thời gian thực.'''

content_en = '''Abstract – In the digital era, the application of artificial intelligence in education is opening up new opportunities to improve teaching quality and classroom management. With the goal of enhancing the effectiveness of classroom management, this study proposes the development of a "Classroom Management Support Application" aimed at automating several processes and improving administrative efficiency. The application’s core feature is attendance tracking through facial recognition, utilizing an MTCNN algorithm combined with MobileFaceNet, achieving 99.55% accuracy on the LFW dataset. The model’s threshold has been fine-tuned to suit real-time facial recognition under various environmental conditions, with an average processing time of 0.01 seconds on a Core i7-10750H CPU, significantly reducing errors and time consumption compared to traditional attendance methods. In addition, the application integrates the YOLO v11 model for person detection, enabling lecturers to monitor and determine the number of students present in the classroom in real-time via surveillance cameras. This feature reaches approximately 85% accuracy on the mAP@50-95 scale. A web application is also provided, allowing lecturers to visualize attendance reports  and statistics, monitor classrooms, and track the attendance process in real-time.'''

references = [
    'F. Schroff, D. Kalenichenko and J. Philbin, "FaceNet: A unified embedding for face recognition and clustering," 2015 IEEE CVPR, 2015.',
    'X. Wu, R. He, Z. Sun and T. Tan, "A Light CNN for Deep Face Representation With Noisy Labels," IEEE TIFS, 2018.',
    'S. Chen et al., MobileFaceNets: Efficient CNNs for Accurate Real-Time Face Verification on Mobile Devices. 2018.',
    'G. Vasavi et al., "People Counting Based on YOLO," 2023 GCAT.',
    'Q. Zhu et al., "Fast Human Detection Using a Cascade of HOGs," CVPR 2006.',
    'R. Khanam and M. Hussain, YOLOv11: An Overview, 2024.'
]

# Main assembly

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('XÂY DỰNG ỨNG DỤNG HỖ TRỢ QUẢN LÝ LỚP HỌC')
run.bold = True
run.font.size = Pt(14)

# Subtitle / authors placeholder
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.add_run('Báo cáo nghiên cứu')

doc.add_paragraph('')

# Table of Contents placeholder
doc.add_paragraph('Mục lục (Cập nhật tự động trong Word):')
doc.add_paragraph('<<Bảng mục lục - update field in Word>>')

# Abstracts
doc.add_heading('Tóm tắt', level=1)
doc.add_paragraph(content_vn)

doc.add_heading('Abstract', level=1)
doc.add_paragraph(content_en)

# Keywords
doc.add_paragraph('Từ Khóa: MTCNN, MobileFaceNet, YOLO, xử lý ảnh, điểm danh qua camera, hỗ trợ quản lý lớp học.')
doc.add_paragraph('Keywords: MTCNN, MobileFaceNet, YOLO, image processing, attendance via camera, classroom management support system.')

# Sections (I, II, III...)

doc.add_heading('I. GIỚI THIỆU', level=1)
doc.add_paragraph('Trong thời đại 4.0, khi dữ liệu được số hóa, các quy trình trở nên tự động hóa, và trí tuệ nhân tạo (AI) được sử dụng rộng rãi. Việc áp dụng những phương pháp trên giúp con người giải phóng sức lao động và giải quyết những hạn chế trong nhiều công việc thủ công còn tồn đọng, giúp bắt kịp với xu hướng thời đại.\n\n' +
                  'Một trong lĩnh vực đang được chú trọng áp dụng các phương pháp này là giáo dục, việc ứng dụng số hóa dữ liệu, tự động hóa và AI vào giáo dục đang mở ra những cơ hội mới để cải thiện chất lượng giảng dạy và quản lý lớp học. Trong nghiên cứu này chúng tôi đề xuất xây dựng một ứng dụng hỗ trợ quản lý lớp học sử dụng AI để tự động hóa quy trình điểm danh, và giám sát thống kê số lượng sinh viên trong lớp theo thời gian thực, trực quan hóa số liệu điểm danh thông qua ứng dụng web, bằng biểu đồ, bảng biểu.')

# Model proposal

doc.add_heading('II. MÔ HÌNH ĐỀ XUẤT', level=1)

doc.add_heading('A. Mô hình nhận diện người', level=2)
doc.add_paragraph('Chúng tôi sử dụng mô hình YOLO phiên bản 11 để huấn luyện mô hình nhận diện người...')

doc.add_paragraph('Dữ liệu huấn luyện: 5300 hình ảnh 1920x1080 với 131870 nhãn. Chia: train 80%, val 15%, test 5%.')

# Table: Comparison (BẢNG 1)

doc.add_paragraph('BẢNG 1. SO SÁNH MÔ HÌNH SAU KHI HUẤN LUYỆN')
rows = [['Mô hình','Params (m)','Time (minute)','Train/box loss','Val/box loss','Train/cls_loss','Val/cls_loss','Precision (B)','Recall(B)','mAP50-95(B)'],
        ['YOLOv11n','2.60','30.00','0.82','0.82','0.46','0.46','0.81','0.87','0.60'],
        ['YOLOv11s','9.4','34.50','0.77','0.77','0.44','0.44','0.91','0.97','0.71'],
        ['YOLOv11m','20.10','57.80','0.76','0.76','0.42','0.42','0.91','0.97','0.71']]

table = doc.add_table(rows=1, cols=len(rows[0]))
hdr_cells = table.rows[0].cells
for i, h in enumerate(rows[0]):
    hdr_cells[i].text = h
for row_data in rows[1:]:
    row_cells = table.add_row().cells
    for i, item in enumerate(row_data):
        row_cells[i].text = item

# Continue with face recognition section

doc.add_heading('B. Mô hình nhận diện khuôn mặt', level=2)
doc.add_paragraph('Sử dụng kết hợp MTCNN (phát hiện) và MobileFaceNet (nhận diện). Dữ liệu thu thập: 71 video sinh viên, xử lý cắt mặt, chuẩn hóa 112x112, augmentation mức độ 1-3.')

doc.add_paragraph('Lựa chọn threshold: chạy từ 0.6 tới 1.6 bước 0.05; lựa chọn threshold 1.2 cho dữ liệu level 2; thời gian xử lý khoảng 10 ms trên CPU Core i7 10750H.')

doc.add_paragraph('Hình 3: Chuẩn hóa kích thước và áp dụng data augmentation theo độ khó tăng dần.\nHình 4: Biểu đồ so sánh độ chính xác với threshold.\nHình 5: Biểu đồ thời gian xử lý ảnh tại threshold 1.2.')

# System architecture and app

doc.add_heading('III. XÂY DỰNG ỨNG DỤNG', level=1)

doc.add_heading('A. Kiến trúc hệ thống', level=2)
doc.add_paragraph('Hệ thống sử dụng kiến trúc microservices gồm: Facial Recognition Services, Human Detection and Counting Services, Auth Service (JWT), Classroom Management Service, Image Handler Services.')

doc.add_heading('B. Xây dựng API', level=2)
doc.add_paragraph('Xác thực người dùng: REST API, Golang, JWT, bcrypt.\nQuản lý lớp học: REST API, Golang.\nDịch vụ xử lý ảnh: NodeJS.\nPhát video trực tiếp: Websocket + Python cho xử lý thời gian thực.')

doc.add_heading('C. Thiết bị điểm danh khuôn mặt', level=2)
doc.add_paragraph('Thiết bị: ESP32-CAM, Facial recognition services, ESP32-S3 + ST7789 hiển thị. Tất cả thiết bị cùng mạng WiFi để trao đổi hình ảnh và socket.')

doc.add_heading('D. Ứng dụng web', level=2)
doc.add_paragraph('React + TypeScript + Tailwind + Vite. 7 chức năng chính: Lựa chọn lớp học, Quản lý điểm danh, Quản lý lớp học, Quản lý sinh viên, Giám sát điểm danh, Giám sát lớp học, Báo cáo thống kê.')

# Results & conclusion

doc.add_heading('IV. KẾT LUẬN', level=1)
doc.add_paragraph('Hệ thống đề xuất tích hợp MTCNN + MobileFaceNet cho điểm danh khuôn mặt (độ chính xác 99.55% trên LFW), YOLOv11s cho phát hiện người (mAP50-95 ~85%). Hệ thống hỗ trợ trực quan hóa, giảm thiểu sai sót điểm danh và cải thiện quản lý lớp học.')

doc.add_heading('Tài liệu tham khảo', level=1)
for ref in references:
    doc.add_paragraph(ref, style='List Number')

# Save
outfile = 'BaoCao_QuanLy_LopHoc.docx'
doc.save(outfile)
print(f'Generated {outfile}')
